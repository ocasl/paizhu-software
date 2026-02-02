# -*- coding: utf-8 -*-
from docx import Document
import os

path = r"e:\CODE\paizhu-software\backend\muban\template_temp.docx"

print(f"处理文件: {path}")

doc = Document(path)

print("=== 正在处理模板文件 ===")

count = 0

# 遍历所有表格单元格
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text.strip()
                    # 检查是否是纯数字 1-12
                    if text in ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12']:
                        new_text = '{' + text + '}'
                        run.text = new_text
                        print(f"替换: '{text}' -> '{new_text}' 位置: 表{t_idx} 行{r_idx} 列{c_idx}")
                        count += 1

# 保存回原文件
output_path = r"e:\CODE\paizhu-software\backend\muban\派驻检察工作日志.docx"
doc.save(output_path)
print(f"\n完成！共替换了 {count} 处，保存到: {output_path}")
