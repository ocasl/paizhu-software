"""
将模板中的***替换为{1}, {2}, {3}...数字占位符
"""
from docx import Document
import sys
import os

def number_placeholders(doc_path, output_path):
    """将***替换为数字占位符"""
    doc = Document(doc_path)
    
    placeholder_count = 0
    replacements = []
    
    # 处理段落
    for para_idx, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        if '***' in original_text:
            for run in paragraph.runs:
                while '***' in run.text:
                    placeholder_count += 1
                    # 记录替换信息
                    context = run.text[:60] if len(run.text) <= 60 else run.text[:57] + "..."
                    replacements.append(f"{placeholder_count}. {context}")
                    run.text = run.text.replace('***', f'{{{placeholder_count}}}', 1)
    
    # 处理表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if '***' in paragraph.text:
                        for run in paragraph.runs:
                            while '***' in run.text:
                                placeholder_count += 1
                                context = run.text[:60] if len(run.text) <= 60 else run.text[:57] + "..."
                                replacements.append(f"{placeholder_count}. [表格{table_idx}] {context}")
                                run.text = run.text.replace('***', f'{{{placeholder_count}}}', 1)
    
    # 保存文件
    doc.save(output_path)
    
    # 打印完整列表
    print("\n替换详情:")
    print("=" * 80)
    for repl in replacements:
        print(repl)
    
    print("\n" + "=" * 80)
    print(f"✅ 共替换了 {placeholder_count} 个占位符")
    print(f"✅ 预期应该有59个占位符")
    
    if placeholder_count == 59:
        print("✅ 数量正确!")
    else:
        print(f"⚠️  警告: 实际{placeholder_count}个，预期59个，相差{59-placeholder_count}个")
    
    print(f"✅ 新文件已保存: {output_path}")
    print("=" * 80)

if __name__ == '__main__':
    template_dir = r"E:\CODE\paizhu-software\backend\muban"
    input_file = os.path.join(template_dir, "派驻检察室月度工作情况报告.docx")
    output_file = os.path.join(template_dir, "派驻检察室月度工作情况报告_numbered.docx")
    
    if not os.path.exists(input_file):
        print(f"❌ 找不到文件: {input_file}")
        sys.exit(1)
    
    print("="*80)
    print("将***替换为数字占位符 {1}, {2}, {3}...")
    print("="*80)
    number_placeholders(input_file, output_file)
