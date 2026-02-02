#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
给派驻检察工作报告事项清单Word模板添加占位符

使用方法：
1. 安装依赖: pip install python-docx
2. 运行脚本: python add-placeholders-to-checklist.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 文件路径
TEMPLATE_PATH = '../muban/派驻检察工作报告事项清单.docx'
OUTPUT_PATH = '../muban/派驻检察工作报告事项清单_with_placeholders.docx'

def add_placeholders():
    """给Word模板添加占位符"""
    
    print('正在读取模板文件...')
    doc = Document(TEMPLATE_PATH)
    
    # 1. 处理标题和基本信息
    print('添加基本信息占位符...')
    
    # 遍历所有段落，找到需要替换的文本
    for para in doc.paragraphs:
        if '派驻监所：' in para.text:
            # 在"派驻监所："后面添加占位符
            para.text = para.text.replace('派驻监所：', '派驻监所：{{prison_name}}')
        
        if '年' in para.text and '月' in para.text:
            # 替换年月占位符
            # 假设格式是 "____年____月"
            para.text = para.text.replace('____年____月', '{{year}}年{{month}}月')
    
    # 2. 处理表格
    print('添加表格占位符...')
    
    # 找到主表格（通常是第一个或第二个表格）
    if len(doc.tables) > 0:
        table = doc.tables[0]  # 第一个表格
        
        print(f'找到表格，共 {len(table.rows)} 行')
        
        # 遍历表格行（跳过表头）
        # 假设第1行是表头，从第2行开始是数据
        for i, row in enumerate(table.rows):
            if i == 0:  # 跳过表头
                continue
            
            # 获取单元格
            cells = row.cells
            
            if len(cells) >= 5:  # 确保有足够的列
                # 序号列（第0列）已经有固定值
                # 报告事项列（第1列）已经有固定值
                # 检察时间列（第2列）已经有固定值
                
                # 报告内容列（第3列）- 添加占位符
                content_cell = cells[3]
                content_cell.text = f'{{{{item{i}_content}}}}'
                
                # 检察情况列（第4列）- 添加占位符
                status_cell = cells[4]
                status_cell.text = f'{{{{item{i}_status}}}}'
                
                print(f'  第{i}行: 添加了 item{i}_content 和 item{i}_status 占位符')
    
    # 3. 保存新文件
    print(f'\n保存模板到: {OUTPUT_PATH}')
    doc.save(OUTPUT_PATH)
    print('✓ 完成！')
    
    # 4. 生成占位符说明文档
    print('\n生成的占位符列表：')
    print('基本信息:')
    print('  {{prison_name}} - 派驻监所名称')
    print('  {{year}} - 年份')
    print('  {{month}} - 月份')
    print('\n表格数据（16行）:')
    for i in range(1, 17):
        print(f'  {{{{item{i}_content}}}} - 第{i}项报告内容')
        print(f'  {{{{item{i}_status}}}} - 第{i}项检察情况')

if __name__ == '__main__':
    try:
        # 检查文件是否存在
        if not os.path.exists(TEMPLATE_PATH):
            print(f'错误: 找不到模板文件 {TEMPLATE_PATH}')
            print('请确保文件路径正确')
            exit(1)
        
        add_placeholders()
        
        print('\n' + '='*50)
        print('下一步：')
        print('1. 检查生成的文件是否正确')
        print('2. 如果正确，将其重命名为原文件名')
        print('3. 修改 Node.js 代码使用新的占位符')
        print('='*50)
        
    except Exception as e:
        print(f'\n错误: {e}')
        import traceback
        traceback.print_exc()
