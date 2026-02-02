"""
分析犯情动态模板，提取罪犯分类字段
"""
from docx import Document
import os

def analyze_prisoner_template(file_path):
    """分析犯情动态文档"""
    print(f"分析文件: {os.path.basename(file_path)}")
    print("=" * 80)
    
    try:
        doc = Document(file_path)
        
        # 显示所有段落
        print("\n📄 文档内容:")
        print("-" * 80)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[{i}] {para.text}")
        
        # 显示表格
        print("\n📊 表格内容:")
        print("-" * 80)
        for t_idx, table in enumerate(doc.tables):
            print(f"\n表格 {t_idx}:")
            for r_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"  行{r_idx}: {' | '.join(row_data)}")
        
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    template_path = r"E:\CODE\paizhu-software\muban\XX省XX监狱2025年某月犯情动态.docx"
    
    if os.path.exists(template_path):
        analyze_prisoner_template(template_path)
    else:
        print(f"❌ 文件不存在: {template_path}")
        print("\n请确认文件路径是否正确")
