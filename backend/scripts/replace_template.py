"""
直接替换Word模板中的***占位符
"""
from docx import Document
import sys
import json

def replace_stars_in_document(doc_path, replacements):
    """在Word文档中替换***占位符"""
    doc = Document(doc_path)
    
    # 替换段落
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if '***' in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace('***', str(value), 1) if '***' in run.text else run.text
   
    # 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if '***' in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace('***', str(value), 1) if '***' in run.text else run.text
    
    return doc

def main():
    if len(sys.argv) != 3:
        print("Usage: python replace_template.py <template_path> <data_json>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    data = json.loads(sys.argv[2])
    
    doc = replace_stars_in_document(template_path, data)
    
    # 输出到stdout (binary)
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    sys.stdout.buffer.write(buffer.getvalue())

if __name__ == '__main__':
    main()
