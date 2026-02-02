"""
分析Word模板内容，查看所有的***占位符
"""
from docx import Document
import os

def analyze_template(file_path):
    """分析模板文件，列出所有包含***的文本"""
    print(f"\n{'='*80}")
    print(f"分析文件: {os.path.basename(file_path)}")
    print(f"{'='*80}\n")
    
    try:
        doc = Document(file_path)
        
        # 收集所有包含***的文本
        paragraphs_with_stars = []
        table_cells_with_stars = []
        
        # 分析段落
        print("📄 段落中的***占位符:")
        print("-" * 80)
        for i, para in enumerate(doc.paragraphs):
            if '***' in para.text:
                paragraphs_with_stars.append((i, para.text))
                print(f"[段落 {i}] {para.text[:100]}")
        
        # 分析表格
        print(f"\n📊 表格中的***占位符:")
        print("-" * 80)
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    if '***' in cell_text:
                        table_cells_with_stars.append((table_idx, row_idx, cell_idx, cell_text))
                        print(f"[表格 {table_idx}, 行 {row_idx}, 列 {cell_idx}]")
                        print(f"  {cell_text[:100]}")
        
        # 统计
        print(f"\n📈 统计:")
        print(f"- 包含***的段落: {len(paragraphs_with_stars)} 个")
        print(f"- 包含***的表格单元格: {len(table_cells_with_stars)} 个")
        
        # 提取所有唯一的***模式
        print(f"\n🔍 检测到的***占位符模式:")
        print("-" * 80)
        patterns = set()
        
        for _, text in paragraphs_with_stars:
            # 简单查找包含***的词组
            words = text.split()
            for word in words:
                if '***' in word:
                    patterns.add(word)
        
        for _, _, _, text in table_cells_with_stars:
            words = text.split()
            for word in words:
                if '***' in word:
                    patterns.add(word)
        
        for pattern in sorted(patterns):
            print(f"  • {pattern}")
        
        print(f"\n共找到 {len(patterns)} 种不同的***模式\n")
        
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()

def main():
    template_dir = r"E:\CODE\paizhu-software\backend\muban"
    
    # 分析两个模板
    files = [
        "派驻检察室月度工作情况报告.docx",
        "派驻检察工作报告事项清单.docx"
    ]
    
    for filename in files:
        file_path = os.path.join(template_dir, filename)
        if os.path.exists(file_path):
            analyze_template(file_path)
        else:
            print(f"❌ 文件不存在: {file_path}")

if __name__ == "__main__":
    main()
