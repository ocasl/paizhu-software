#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从头创建事项清单模板
确保占位符格式正确，不会被Word拆分
"""

import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_checklist_template():
    """创建事项清单模板"""
    
    # 输出文件路径
    template_dir = os.path.join(os.path.dirname(__file__), '..', 'muban')
    output_file = os.path.join(template_dir, '派驻检察工作报告事项清单_new.docx')
    
    print('开始创建事项清单模板...\n')
    
    try:
        # 创建新文档
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)
        
        # 添加标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('江西省南昌长塘地区人民检察院\n派驻检察工作报告事项清单')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加派驻监所和年月信息
        info_para = doc.add_paragraph()
        run1 = info_para.add_run('派驻监所: ')
        run1.font.size = Pt(12)
        run2 = info_para.add_run('{{prison_name}}')
        run2.font.size = Pt(12)
        run2.font.bold = True
        
        run3 = info_para.add_run('        ')
        run3.font.size = Pt(12)
        
        run4 = info_para.add_run('{{year}}')
        run4.font.size = Pt(12)
        run4.font.bold = True
        run5 = info_para.add_run(' 年 ')
        run5.font.size = Pt(12)
        run6 = info_para.add_run('{{month}}')
        run6.font.size = Pt(12)
        run6.font.bold = True
        run7 = info_para.add_run(' 月')
        run7.font.size = Pt(12)
        
        # 添加空行
        doc.add_paragraph()
        
        # 创建表格 (4列: 序号、报告事项、检察时间、报告内容、检察情况)
        table = doc.add_table(rows=17, cols=4)  # 1个表头 + 16个事项
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '序号'
        header_cells[1].text = '报告事项'
        header_cells[2].text = '检察时间'
        header_cells[3].text = '检察情况'
        
        # 设置表头格式
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(12)
        
        # 16项标准事项
        items = [
            ('1', '监狱发生罪犯脱逃、自伤自残、自杀死亡、重大疫情、重大生产安全事故的情况报告', '及时'),
            ('2', '罪犯死亡事件调查及处理报告', '及时'),
            ('3', '监狱开展重大监管改造业务活动的情况报告', '及时'),
            ('4', '监狱民警受到党纪行政处罚情况', '及时'),
            ('5', '监狱新任职领导情况列表', '及时'),
            ('6', '监狱提请罪犯减刑、假释、暂予监外执行花名册', '每批次'),
            ('7', '抽查重点时段、重点环节监控录像发现的情况', '每日'),
            ('8', '对监狱医院禁闭室检察情况', '每周'),
            ('9', '罪犯外伤检察', '每周'),
            ('10', '对刑释前罪犯和新入监罪犯谈话情况', '每周'),
            ('11', '开启检察官信箱或检察中发现具有价值的案件线索', '每周'),
            ('12', '检查发现罪犯私藏使用违禁品的情况', '每周'),
            ('13', '对监狱会见场所检察情况', '每月'),
            ('14', '参加监狱犯情分析会情况', '每月'),
            ('15', '记过以上处分的监督情况', '每月'),
            ('16', '狱内勤杂岗位和辅助生产岗位罪犯每月增减情况', '每月')
        ]
        
        # 填充表格内容
        for i, (num, content, frequency) in enumerate(items, start=1):
            row = table.rows[i]
            
            # 序号
            row.cells[0].text = num
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 报告事项 - 使用占位符
            cell = row.cells[1]
            cell.text = ''  # 清空
            para = cell.paragraphs[0]
            run = para.add_run(f'{{{{item{i}_content}}}}')
            run.font.size = Pt(11)
            
            # 检察时间
            row.cells[2].text = frequency
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 检察情况 - 使用占位符
            cell = row.cells[3]
            cell.text = ''  # 清空
            para = cell.paragraphs[0]
            run = para.add_run(f'{{{{item{i}_status}}}}')
            run.font.size = Pt(11)
        
        # 设置列宽
        for row in table.rows:
            row.cells[0].width = Inches(0.5)   # 序号
            row.cells[1].width = Inches(3.5)   # 报告事项
            row.cells[2].width = Inches(0.8)   # 检察时间
            row.cells[3].width = Inches(2.0)   # 检察情况
        
        # 保存文档
        doc.save(output_file)
        
        print(f'✅ 模板创建成功！')
        print(f'文件保存到: {output_file}\n')
        print('占位符列表:')
        print('  - {{prison_name}} : 派驻监所名称')
        print('  - {{year}} : 年份')
        print('  - {{month}} : 月份')
        for i in range(1, 17):
            print(f'  - {{{{item{i}_content}}}} : 第{i}项报告事项内容')
            print(f'  - {{{{item{i}_status}}}} : 第{i}项检察情况')
        
        print('\n下一步:')
        print('1. 打开生成的文件检查格式')
        print('2. 如果正确，将其重命名为: 派驻检察工作报告事项清单_with_placeholders.docx')
        print('3. 运行测试: node backend/scripts/test-checklist-generation.js')
        
        return True
        
    except Exception as e:
        print(f'\n❌ 创建失败: {str(e)}')
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    success = create_checklist_template()
    sys.exit(0 if success else 1)
