#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
在事项清单模板中插入占位符
使模板可以被docxtemplater正确填充
"""

from docx import Document
from pathlib import Path
import sys

def insert_placeholders():
    """在模板中插入占位符"""
    
    # 读取模板
    template_path = Path(__file__).parent.parent / 'muban' / '派驻检察工作报告事项清单_modified.docx'
    output_path = Path(__file__).parent.parent / 'muban' / '派驻检察工作报告事项清单_with_placeholders.docx'
    
    print(f'读取模板: {template_path}')
    
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f'错误: 无法读取模板文件 - {e}')
        return False
    
    print(f'模板加载成功，共 {len(doc.paragraphs)} 个段落')
    
    # 1. 在标题段落中插入占位符
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # 查找包含"派驻监所："的段落
        if '派驻监所：' in text:
            print(f'\n找到标题段落 (段落 {i}): "{text}"')
            
            # 清空段落
            for run in para.runs:
                run.text = ''
            
            # 重新构建段落，插入占位符
            para.add_run('派驻监所：')
            para.add_run('{prison_name}')
            para.add_run('                              ')
            para.add_run('{year}')
            para.add_run('年')
            para.add_run('    ')
            para.add_run('{month}')
            para.add_run('月')
            
            print(f'  ✓ 已插入占位符: {{prison_name}}, {{year}}, {{month}}')
            break
    
    # 2. 在表格的第5列（检察情况）插入占位符
    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f'\n找到表格，共 {len(table.rows)} 行')
        
        # 跳过表头（第1行），从第2行开始（对应16项事项）
        for row_idx in range(1, min(17, len(table.rows))):
            row = table.rows[row_idx]
            
            if len(row.cells) >= 5:
                # 第5列是检察情况（索引4）
                cell = row.cells[4]
                
                # 清空单元格内容
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                
                # 插入占位符 {status1}, {status2}, ... {status16}
                placeholder = f'{{status{row_idx}}}'
                
                # 确保至少有一个段落
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                
                cell.paragraphs[0].add_run(placeholder)
                
                if row_idx <= 3:  # 只打印前3行
                    print(f'  行 {row_idx}: 插入占位符 {placeholder}')
        
        print(f'  ✓ 已为16项事项插入占位符 {{status1}} - {{status16}}')
    
    # 3. 在表格的第4列（报告内容）也插入占位符（如果需要）
    if len(doc.tables) > 0:
        table = doc.tables[0]
        
        for row_idx in range(1, min(17, len(table.rows))):
            row = table.rows[row_idx]
            
            if len(row.cells) >= 5:
                # 第4列是报告内容（索引3）
                cell = row.cells[3]
                
                # 清空单元格内容
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                
                # 插入占位符 {content1}, {content2}, ... {content16}
                placeholder = f'{{content{row_idx}}}'
                
                # 确保至少有一个段落
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                
                cell.paragraphs[0].add_run(placeholder)
        
        print(f'  ✓ 已为16项事项插入占位符 {{content1}} - {{content16}}')
    
    # 保存新模板
    print(f'\n保存新模板到: {output_path}')
    doc.save(output_path)
    
    print('\n✅ 占位符插入成功！')
    print(f'\n新模板已保存: {output_path.name}')
    print('\n占位符列表:')
    print('  - {prison_name}: 监狱名称')
    print('  - {year}: 年份')
    print('  - {month}: 月份')
    print('  - {status1} - {status16}: 16项检察情况')
    print('  - {content1} - {content16}: 16项报告内容')
    
    return True

if __name__ == '__main__':
    success = insert_placeholders()
    sys.exit(0 if success else 1)
