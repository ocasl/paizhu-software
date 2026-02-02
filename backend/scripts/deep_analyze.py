"""
深度分析Word模板 - 显示完整文本内容
"""
from docx import Document
import os

def deep_analyze(file_path):
    """深度分析模板"""
    print(f"\n{'='*80}")
    print(f"深度分析: {os.path.basename(file_path)}")
    print(f"{'='*80}\n")
    
    try:
        doc = Document(file_path)
        
        print("📄 所有段落内容:")
        print("-" * 80)
        for i, para in enumerate(doc.paragraphs[:50]):  # 只显示前50个段落
            if para.text.strip():
                print(f"[{i}] {para.text}")
        
        print(f"\n📊 所有表格内容:")
        print("-" * 80)
        for t_idx, table in enumerate(doc.tables):
            print(f"\n表格 {t_idx}:")
            for r_idx, row in enumerate(table.rows[:10]):  # 只显示前10行
                row_text = " | ".join([cell.text.strip()[:30] for cell in row.cells])
                if row_text.strip():
                    print(f"  行{r_idx}: {row_text}")
        
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()

def main():
    template_dir = r"E:\CODE\paizhu-software\backend\muban"
    
    # 检查所有docx文件
    print("检查muban目录下的所有.docx文件:")
    for filename in os.listdir(template_dir):
        if filename.endswith('.docx') and not filename.startswith('~'):
            print(f"\n找到: {filename}")
    
    # 分析原始文件
    files = [
        "派驻检察室月度工作情况报告.docx",
        "派驻检察工作报告事项清单.docx"
    ]
    
    for filename in files:
        file_path = os.path.join(template_dir, filename)
        if os.path.exists(file_path):
            deep_analyze(file_path)

if __name__ == "__main__":
    main()
