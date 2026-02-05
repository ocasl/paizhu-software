#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试脚本：查看Word文档的原始文本内容
"""

import sys
from docx import Document

if len(sys.argv) < 2:
    print("用法: python test_parse.py <docx文件路径>")
    sys.exit(1)

docx_path = sys.argv[1]

try:
    doc = Document(docx_path)
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    
    print("=" * 80)
    print("Word文档完整文本内容:")
    print("=" * 80)
    print(full_text)
    print("=" * 80)
    print(f"\n总字符数: {len(full_text)}")
    print(f"总段落数: {len(doc.paragraphs)}")
    
    # 查找关键段落
    print("\n" + "=" * 80)
    print("包含'在押罪犯'的段落:")
    print("=" * 80)
    for i, para in enumerate(doc.paragraphs):
        if '在押' in para.text or '罪犯' in para.text:
            print(f"段落 {i}: {para.text}")
    
except Exception as e:
    print(f"错误: {e}")
